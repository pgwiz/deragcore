"""File parsing - Extract text from PDF and DOCX files."""

import logging
from typing import Tuple, Dict, Any, Optional
import io

logger = logging.getLogger(__name__)


class FileParser:
    """Parse PDF and DOCX files to extract clean text."""

    @staticmethod
    def parse_pdf(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF using PyMuPDF.

        Args:
            file_bytes: Raw PDF file bytes

        Returns:
            (text: str, metadata: Dict) where metadata contains page_count, etc.

        Raises:
            ValueError: If PDF is corrupted or encrypted
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF (fitz) required for PDF parsing")

        try:
            pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
            text_parts = []
            page_count = len(pdf_document)

            for page_num, page in enumerate(pdf_document):
                try:
                    text = page.get_text()
                    if text.strip():
                        # Add page marker for metadata
                        text_parts.append(f"\n--- Page {page_num + 1} ---\n{text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page {page_num + 1}: {e}")
                    continue

            pdf_document.close()

            full_text = "\n".join(text_parts)

            if not full_text.strip():
                raise ValueError("PDF contains no extractable text")

            metadata = {
                "page_count": page_count,
                "file_type": "pdf",
                "char_count": len(full_text),
            }

            logger.info(f"Parsed PDF: {page_count} pages, {len(full_text)} characters")
            return full_text, metadata

        except Exception as e:
            error_msg = f"PDF parsing error: {str(e)}"
            logger.error(error_msg)
            raise ValueError(error_msg) from e

    @staticmethod
    def parse_docx(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from DOCX using python-docx.

        Args:
            file_bytes: Raw DOCX file bytes

        Returns:
            (text: str, metadata: Dict) with paragraph_count, etc.

        Raises:
            ValueError: If DOCX is corrupted or invalid
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required for DOCX parsing")

        try:
            # Load DOCX from bytes
            doc_stream = io.BytesIO(file_bytes)
            doc = Document(doc_stream)

            text_parts = []
            for para_num, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)

            full_text = "\n".join(text_parts)

            if not full_text.strip():
                raise ValueError("DOCX contains no extractable text")

            # Extract metadata
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "file_type": "docx",
                "char_count": len(full_text),
            }

            # Try to get core properties (title, author, etc.)
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
            except Exception:
                pass  # Core properties might not be available

            logger.info(f"Parsed DOCX: {len(doc.paragraphs)} paragraphs, {len(full_text)} characters")
            return full_text, metadata

        except Exception as e:
            error_msg = f"DOCX parsing error: {str(e)}"
            logger.error(error_msg)
            raise ValueError(error_msg) from e

    @staticmethod
    def parse(file_bytes: bytes, content_type: str) -> Tuple[str, Dict[str, Any]]:
        """
        Parse file based on content type.

        Args:
            file_bytes: Raw file bytes
            content_type: MIME type (application/pdf or application/vnd.openxmlformats-...)

        Returns:
            (text, metadata) tuple

        Raises:
            ValueError: If content type not supported or file is invalid
        """
        if "pdf" in content_type.lower():
            return FileParser.parse_pdf(file_bytes)
        elif "wordprocessingml" in content_type.lower() or "word" in content_type.lower():
            return FileParser.parse_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {content_type}")
